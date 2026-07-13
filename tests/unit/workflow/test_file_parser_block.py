"""
Tests for FileParserBlock DOCX support.

Covers file type detection, validation, text extraction (paragraphs + tables),
token truncation, and error handling for DOCX files.
"""

from __future__ import annotations

import asyncio
import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from unittest.mock import AsyncMock, MagicMock

import docx
import pytest

from skyvern.forge.sdk.api.llm.exceptions import InvalidLLMResponseFormat
from skyvern.forge.sdk.workflow.exceptions import InvalidFileType
from skyvern.forge.sdk.workflow.models.block import BlockType, FileParserBlock, PDFParserBlock
from skyvern.forge.sdk.workflow.models.parameter import OutputParameter, ParameterType
from skyvern.schemas.workflows import BlockResult, BlockStatus, FileType


def _make_output_parameter(key: str) -> OutputParameter:
    return OutputParameter(
        parameter_type=ParameterType.OUTPUT,
        key=key,
        description="test",
        output_parameter_id="test-output-id",
        workflow_id="test-workflow-id",
        created_at=datetime.now(timezone.utc),
        modified_at=datetime.now(timezone.utc),
    )


def _make_file_parser_block(file_url: str, file_type: FileType) -> FileParserBlock:
    return FileParserBlock(
        label="test_file_parser",
        block_type=BlockType.FILE_URL_PARSER,
        output_parameter=_make_output_parameter("test_output"),
        file_url=file_url,
        file_type=file_type,
    )


def _make_pdf_parser_block(file_url: str) -> PDFParserBlock:
    return PDFParserBlock(
        label="test_pdf_parser",
        block_type=BlockType.PDF_PARSER,
        output_parameter=_make_output_parameter("test_output"),
        file_url=file_url,
    )


def _create_docx(
    path: Path,
    paragraphs: list[str] | None = None,
    table_rows: list[list[str]] | None = None,
) -> Path:
    """Create a DOCX file with optional paragraphs and tables."""
    doc = docx.Document()
    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)
    if table_rows:
        cols = len(table_rows[0])
        table = doc.add_table(rows=len(table_rows), cols=cols)
        for i, row_data in enumerate(table_rows):
            for j, cell_text in enumerate(row_data):
                table.rows[i].cells[j].text = cell_text
    doc.save(str(path))
    return path


class TestDetectFileTypeFromUrl:
    """Tests for _detect_file_type_from_url with DOCX extensions."""

    def _detect(self, url: str, file_path: str | None = None) -> FileType:
        block = _make_file_parser_block(url, FileType.CSV)
        return block._detect_file_type_from_url(url, file_path=file_path)

    def test_docx_extension(self) -> None:
        assert self._detect("https://example.com/file.docx") == FileType.DOCX

    def test_doc_extension_raises_error(self) -> None:
        # Legacy .doc (Word 97-2003) is not supported by python-docx
        with pytest.raises(InvalidFileType, match="Legacy .doc format"):
            self._detect("https://example.com/file.doc")

    def test_docx_with_query_params(self) -> None:
        assert self._detect("https://example.com/file.docx?token=abc&v=1") == FileType.DOCX

    def test_docx_case_insensitive(self) -> None:
        assert self._detect("https://example.com/file.DOCX") == FileType.DOCX

    def test_other_extensions_unchanged(self) -> None:
        assert self._detect("https://example.com/file.pdf") == FileType.PDF
        assert self._detect("https://example.com/file.xlsx") == FileType.EXCEL
        assert self._detect("https://example.com/file.csv") == FileType.CSV
        assert self._detect("https://example.com/file.png") == FileType.IMAGE

    def test_no_extension_without_file_path_falls_back_to_csv(self) -> None:
        assert self._detect("https://example.com/34371136523") == FileType.CSV

    def test_no_extension_with_pdf_file_detected_as_pdf(self, tmp_path: Path) -> None:
        # Create a minimal valid PDF file
        pdf_path = tmp_path / "no_ext_file"
        pdf_path.write_bytes(b"%PDF-1.5\n1 0 obj\n<< /Type /Catalog >>\nendobj\n%%EOF")
        assert self._detect("https://example.com/34371136523", file_path=str(pdf_path)) == FileType.PDF

    def test_no_extension_with_unknown_file_falls_back_to_csv(self, tmp_path: Path) -> None:
        # Plain text file — filetype.guess returns None for text
        txt_path = tmp_path / "unknown_file"
        txt_path.write_text("just,some,csv,data\n1,2,3,4")
        assert self._detect("https://example.com/some_file", file_path=str(txt_path)) == FileType.CSV

    def test_query_params_only_url_with_pdf_file(self, tmp_path: Path) -> None:
        # URL like /download?id=123 — no file extension visible
        pdf_path = tmp_path / "downloaded"
        pdf_path.write_bytes(b"%PDF-1.5\n1 0 obj\n<< /Type /Catalog >>\nendobj\n%%EOF")
        assert self._detect("https://example.com/download?id=123", file_path=str(pdf_path)) == FileType.PDF


class TestValidateFileType:
    """Tests for validate_file_type with DOCX files."""

    def test_valid_docx(self, tmp_path: Path) -> None:
        path = _create_docx(tmp_path / "valid.docx", paragraphs=["Hello"])
        block = _make_file_parser_block("https://example.com/valid.docx", FileType.DOCX)
        # Should not raise
        block.validate_file_type("https://example.com/valid.docx", str(path))

    def test_plain_text_with_docx_extension(self, tmp_path: Path) -> None:
        path = tmp_path / "fake.docx"
        path.write_text("This is plain text, not a DOCX file.")
        block = _make_file_parser_block("https://example.com/fake.docx", FileType.DOCX)
        with pytest.raises(InvalidFileType):
            block.validate_file_type("https://example.com/fake.docx", str(path))

    def test_empty_file(self, tmp_path: Path) -> None:
        path = tmp_path / "empty.docx"
        path.write_bytes(b"")
        block = _make_file_parser_block("https://example.com/empty.docx", FileType.DOCX)
        with pytest.raises(InvalidFileType):
            block.validate_file_type("https://example.com/empty.docx", str(path))


@pytest.mark.asyncio
class TestParseDocxFile:
    """Tests for _parse_docx_file text extraction."""

    async def test_paragraphs_joined_by_newline(self, tmp_path: Path) -> None:
        path = _create_docx(tmp_path / "paras.docx", paragraphs=["Hello", "World"])
        block = _make_file_parser_block("https://example.com/paras.docx", FileType.DOCX)
        result = await block._parse_docx_file(str(path))
        assert result == "Hello\nWorld"

    async def test_empty_paragraphs_skipped(self, tmp_path: Path) -> None:
        path = _create_docx(tmp_path / "blanks.docx", paragraphs=["Hello", "", "   ", "World"])
        block = _make_file_parser_block("https://example.com/blanks.docx", FileType.DOCX)
        result = await block._parse_docx_file(str(path))
        assert result == "Hello\nWorld"

    async def test_table_rows_formatted_with_pipe(self, tmp_path: Path) -> None:
        path = _create_docx(
            tmp_path / "table.docx",
            table_rows=[["Name", "Age"], ["Alice", "30"]],
        )
        block = _make_file_parser_block("https://example.com/table.docx", FileType.DOCX)
        result = await block._parse_docx_file(str(path))
        assert result == "Name | Age\nAlice | 30"

    async def test_mixed_paragraphs_and_tables(self, tmp_path: Path) -> None:
        path = _create_docx(
            tmp_path / "mixed.docx",
            paragraphs=["Intro"],
            table_rows=[["Col1", "Col2"], ["A", "B"]],
        )
        block = _make_file_parser_block("https://example.com/mixed.docx", FileType.DOCX)
        result = await block._parse_docx_file(str(path))
        assert result == "Intro\nCol1 | Col2\nA | B"

    async def test_empty_document(self, tmp_path: Path) -> None:
        path = _create_docx(tmp_path / "empty.docx")
        block = _make_file_parser_block("https://example.com/empty.docx", FileType.DOCX)
        result = await block._parse_docx_file(str(path))
        assert result == ""

    async def test_empty_table_cells_skipped(self, tmp_path: Path) -> None:
        path = _create_docx(
            tmp_path / "sparse.docx",
            table_rows=[["Name", "", "Age"], ["", "", ""]],
        )
        block = _make_file_parser_block("https://example.com/sparse.docx", FileType.DOCX)
        result = await block._parse_docx_file(str(path))
        # First row: "Name" and "Age" (empty cell skipped), second row: all empty -> skipped
        assert result == "Name | Age"

    async def test_multiple_tables(self, tmp_path: Path) -> None:
        doc = docx.Document()
        t1 = doc.add_table(rows=1, cols=2)
        t1.rows[0].cells[0].text = "T1C1"
        t1.rows[0].cells[1].text = "T1C2"
        t2 = doc.add_table(rows=1, cols=2)
        t2.rows[0].cells[0].text = "T2C1"
        t2.rows[0].cells[1].text = "T2C2"
        path = tmp_path / "multi_table.docx"
        doc.save(str(path))

        block = _make_file_parser_block("https://example.com/multi_table.docx", FileType.DOCX)
        result = await block._parse_docx_file(str(path))
        assert result == "T1C1 | T1C2\nT2C1 | T2C2"


@pytest.mark.asyncio
class TestParseDocxFileTokenTruncation:
    """Tests for _parse_docx_file token limit enforcement."""

    async def test_paragraphs_truncated(self, tmp_path: Path) -> None:
        # Create many paragraphs that will exceed a small token limit
        paragraphs = [f"This is paragraph number {i} with some text content." for i in range(100)]
        path = _create_docx(tmp_path / "long.docx", paragraphs=paragraphs)
        block = _make_file_parser_block("https://example.com/long.docx", FileType.DOCX)
        result = await block._parse_docx_file(str(path), max_tokens=20)
        lines = result.split("\n")
        assert len(lines) < len(paragraphs)
        # Each included line should be a valid paragraph
        for line in lines:
            assert line.startswith("This is paragraph number")

    async def test_tables_truncated(self, tmp_path: Path) -> None:
        table_rows = [[f"R{i}C1", f"R{i}C2", f"R{i}C3"] for i in range(100)]
        path = _create_docx(tmp_path / "big_table.docx", table_rows=table_rows)
        block = _make_file_parser_block("https://example.com/big_table.docx", FileType.DOCX)
        result = await block._parse_docx_file(str(path), max_tokens=20)
        lines = result.split("\n")
        assert len(lines) < len(table_rows)

    async def test_tables_skipped_when_paragraphs_exhaust_budget(self, tmp_path: Path) -> None:
        paragraphs = [f"Long paragraph {i} with lots of content to fill tokens." for i in range(100)]
        table_rows = [["Should", "Not", "Appear"]]
        path = _create_docx(tmp_path / "para_heavy.docx", paragraphs=paragraphs, table_rows=table_rows)
        block = _make_file_parser_block("https://example.com/para_heavy.docx", FileType.DOCX)
        result = await block._parse_docx_file(str(path), max_tokens=20)
        assert "Should" not in result
        assert "Not" not in result
        assert "Appear" not in result


@pytest.mark.asyncio
class TestParseDocxFileErrorHandling:
    """Tests for _parse_docx_file error handling."""

    async def test_corrupt_file(self, tmp_path: Path) -> None:
        path = tmp_path / "corrupt.docx"
        path.write_bytes(b"\x00\x01\x02\x03random bytes")
        block = _make_file_parser_block("https://example.com/corrupt.docx", FileType.DOCX)
        with pytest.raises(InvalidFileType):
            await block._parse_docx_file(str(path))

    async def test_nonexistent_file(self, tmp_path: Path) -> None:
        block = _make_file_parser_block("https://example.com/missing.docx", FileType.DOCX)
        with pytest.raises(InvalidFileType):
            await block._parse_docx_file(str(tmp_path / "nonexistent.docx"))


class TestExtractFileUrlFromBlockOutput:
    """Tests for _extract_file_url_from_block_output – unstructured block output parsing."""

    def _extract(self, value: object) -> str | None:
        return FileParserBlock._extract_file_url_from_block_output(value)

    # --- dict inputs ---

    def test_dict_with_downloaded_files_returns_first_url(self) -> None:
        value = {"downloaded_files": [{"url": "https://example.com/file.pdf", "checksum": None}]}
        assert self._extract(value) == "https://example.com/file.pdf"

    def test_dict_multiple_downloaded_files_returns_first(self) -> None:
        value = {
            "downloaded_files": [
                {"url": "https://example.com/first.pdf"},
                {"url": "https://example.com/second.pdf"},
            ]
        }
        assert self._extract(value) == "https://example.com/first.pdf"

    def test_dict_with_extra_fields_still_extracts_url(self) -> None:
        value = {
            "extracted_information": {"key": "value"},
            "downloaded_files": [{"url": "https://s3.amazonaws.com/bucket/report.xlsx", "filename": "report.xlsx"}],
        }
        assert self._extract(value) == "https://s3.amazonaws.com/bucket/report.xlsx"

    def test_dict_empty_downloaded_files_returns_none(self) -> None:
        assert self._extract({"downloaded_files": []}) is None

    def test_dict_missing_downloaded_files_returns_none(self) -> None:
        assert self._extract({"extracted_information": {"foo": "bar"}}) is None

    def test_dict_downloaded_files_item_missing_url_returns_none(self) -> None:
        assert self._extract({"downloaded_files": [{"filename": "file.pdf"}]}) is None

    def test_dict_downloaded_files_item_empty_url_returns_none(self) -> None:
        assert self._extract({"downloaded_files": [{"url": ""}]}) is None

    # --- JSON string inputs ---

    def test_json_string_with_downloaded_files_returns_url(self) -> None:
        value = json.dumps({"downloaded_files": [{"url": "https://example.com/file.csv"}]})
        assert self._extract(value) == "https://example.com/file.csv"

    def test_json_string_without_downloaded_files_returns_none(self) -> None:
        value = json.dumps({"extracted_information": {"k": "v"}})
        assert self._extract(value) is None

    # --- Python dict repr strings (produced by Jinja {{ block_output }} rendering) ---

    def test_python_repr_string_with_downloaded_files_returns_url(self) -> None:
        value = "{'downloaded_files': [{'url': 'https://example.com/report.pdf', 'checksum': None}]}"
        assert self._extract(value) == "https://example.com/report.pdf"

    def test_python_repr_string_without_downloaded_files_returns_none(self) -> None:
        value = "{'extracted_information': {'a': 1}}"
        assert self._extract(value) is None

    # --- Plain URL strings (should not be extracted, returns None) ---

    def test_plain_url_string_returns_none(self) -> None:
        assert self._extract("https://example.com/file.pdf") is None

    def test_plain_string_returns_none(self) -> None:
        assert self._extract("not a url or dict") is None

    # --- Other types ---

    def test_none_returns_none(self) -> None:
        assert self._extract(None) is None

    def test_list_returns_none(self) -> None:
        assert self._extract([{"url": "https://example.com/file.pdf"}]) is None

    def test_integer_returns_none(self) -> None:
        assert self._extract(42) is None


@pytest.mark.asyncio
class TestExtractWithAiSerialization:
    """Tests for _extract_with_ai content serialization."""

    async def test_list_content_serialized_as_compact_json(self) -> None:
        """CSV/Excel data (list[dict]) must use compact JSON to minimize tokens."""
        block = _make_file_parser_block("https://example.com/data.xlsx", FileType.EXCEL)
        block.json_schema = {"type": "object"}
        records = [{"name": "Alice", "age": 30}, {"name": "Bob", "age": 25}]

        with pytest.MonkeyPatch.context() as mp:
            mock_handler = AsyncMock(return_value={})
            mp.setattr(
                "skyvern.forge.sdk.workflow.models.block.LLMAPIHandlerFactory.get_override_llm_api_handler",
                lambda *a, **kw: mock_handler,
            )
            mock_load = MagicMock(return_value="prompt")
            mp.setattr("skyvern.forge.sdk.workflow.models.block.prompt_engine.load_prompt", mock_load)

            await block._extract_with_ai(records, MagicMock())

            _, kwargs = mock_load.call_args
            content_str = kwargs["extracted_text_content"]

            assert content_str == json.dumps(records, separators=(",", ":"))
            assert json.loads(content_str) == records

    async def test_string_content_passed_unchanged(self) -> None:
        """Non-list content (PDF/DOCX text) must pass through unchanged."""
        block = _make_file_parser_block("https://example.com/doc.pdf", FileType.PDF)
        block.json_schema = {"type": "object"}

        with pytest.MonkeyPatch.context() as mp:
            mock_handler = AsyncMock(return_value={})
            mp.setattr(
                "skyvern.forge.sdk.workflow.models.block.LLMAPIHandlerFactory.get_override_llm_api_handler",
                lambda *a, **kw: mock_handler,
            )
            mock_load = MagicMock(return_value="prompt")
            mp.setattr("skyvern.forge.sdk.workflow.models.block.prompt_engine.load_prompt", mock_load)

            await block._extract_with_ai("Hello\nWorld", MagicMock())

            _, kwargs = mock_load.call_args
            assert kwargs["extracted_text_content"] == "Hello\nWorld"


@pytest.mark.asyncio
class TestExtractWithAiSchemaValidation:
    """Tests for schema adherence in FileParserBlock AI extraction."""

    @staticmethod
    def _patch_prompt_and_handler(mp: pytest.MonkeyPatch, handler: AsyncMock, prompt: str = "base prompt") -> None:
        mp.setattr(
            "skyvern.forge.sdk.workflow.models.block.LLMAPIHandlerFactory.get_override_llm_api_handler",
            lambda *a, **kw: handler,
        )
        mp.setattr(
            "skyvern.forge.sdk.workflow.models.block.prompt_engine.load_prompt",
            MagicMock(return_value=prompt),
        )

    async def test_array_schema_does_not_force_dict(self) -> None:
        block = _make_file_parser_block("https://example.com/data.csv", FileType.CSV)
        block.json_schema = {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {"name": {"type": "string"}},
                "required": ["name"],
            },
        }
        captured: dict[str, object] = {}

        async def fake_handler(**kwargs: Any) -> list[dict[str, str]]:
            captured["force_dict"] = kwargs["force_dict"]
            return [{"name": "Alice"}]

        with pytest.MonkeyPatch.context() as mp:
            handler = AsyncMock(side_effect=fake_handler)
            self._patch_prompt_and_handler(mp, handler)

            result = await block._extract_with_ai("name\nAlice", MagicMock())

        assert captured["force_dict"] is False
        assert result == [{"name": "Alice"}]

    async def test_object_schema_wrong_root_retries_without_prevalidation_coercion(self) -> None:
        block = _make_file_parser_block("https://example.com/data.csv", FileType.CSV)
        block.json_schema = {
            "type": "object",
            "properties": {"name": {"type": "string"}},
            "required": ["name"],
        }
        responses: list[Any] = [[{"name": "Alice"}], {"name": "Alice"}]
        prompts: list[str] = []
        force_dict_values: list[bool] = []

        async def fake_handler(**kwargs: Any) -> Any:
            prompts.append(kwargs["prompt"])
            force_dict_values.append(kwargs["force_dict"])
            return responses.pop(0)

        with pytest.MonkeyPatch.context() as mp:
            handler = AsyncMock(side_effect=fake_handler)
            self._patch_prompt_and_handler(mp, handler)

            result = await block._extract_with_ai("name\nAlice", MagicMock())

        assert result == {"name": "Alice"}
        assert handler.await_count == 2
        assert force_dict_values == [False, False]
        assert "previous response failed JSON schema validation" in prompts[1]
        assert "expected type object, got array" in prompts[1]

    async def test_invalid_schema_fails_before_llm_retry(self) -> None:
        block = _make_file_parser_block("https://example.com/data.csv", FileType.CSV)
        block.json_schema = {"type": 123}

        with pytest.MonkeyPatch.context() as mp:
            handler = AsyncMock()
            self._patch_prompt_and_handler(mp, handler)

            with pytest.raises(ValueError, match="File parser JSON schema is invalid"):
                await block._extract_with_ai("name\nAlice", MagicMock())

        handler.assert_not_awaited()

    async def test_schema_validation_failure_retries_with_sanitized_prompt(self) -> None:
        block = _make_file_parser_block("https://example.com/data.csv", FileType.CSV)
        block.json_schema = {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {"name": {"type": "string"}},
                "required": ["name"],
            },
        }
        secret_like_value = "customer-private-value-" + ("x" * 300)
        responses: list[Any] = [{"name": secret_like_value}, [{"name": "Alice"}]]
        prompts: list[str] = []

        async def fake_handler(**kwargs: Any) -> Any:
            prompts.append(kwargs["prompt"])
            return responses.pop(0)

        with pytest.MonkeyPatch.context() as mp:
            handler = AsyncMock(side_effect=fake_handler)
            self._patch_prompt_and_handler(mp, handler)

            result = await block._extract_with_ai("name\nAlice", MagicMock())

        assert result == [{"name": "Alice"}]
        assert handler.await_count == 2
        assert prompts[0] == "base prompt"
        assert "previous response failed JSON schema validation" in prompts[1]
        assert "expected type array, got object" in prompts[1]
        assert secret_like_value not in prompts[1]
        assert "customer-private-value" not in prompts[1]

    async def test_response_format_failure_retries_with_sanitized_prompt(self) -> None:
        block = _make_file_parser_block("https://example.com/data.csv", FileType.CSV)
        block.json_schema = {
            "type": "object",
            "properties": {"name": {"type": "string"}},
            "required": ["name"],
        }
        raw_bad_response = "not-json-customer-private-value-" + ("x" * 300)
        prompts: list[str] = []

        async def fake_handler(**kwargs: Any) -> dict[str, str]:
            prompts.append(kwargs["prompt"])
            if len(prompts) == 1:
                raise InvalidLLMResponseFormat(raw_bad_response)
            return {"name": "Alice"}

        with pytest.MonkeyPatch.context() as mp:
            handler = AsyncMock(side_effect=fake_handler)
            self._patch_prompt_and_handler(mp, handler)

            result = await block._extract_with_ai("name\nAlice", MagicMock())

        assert result == {"name": "Alice"}
        assert handler.await_count == 2
        assert "InvalidLLMResponseFormat" in prompts[1]
        assert raw_bad_response not in prompts[1]
        assert "customer-private-value" not in prompts[1]


@pytest.mark.asyncio
class TestPDFParserSchemaValidation:
    """Tests for schema adherence in the deprecated PDFParserBlock."""

    @staticmethod
    def _patch_execute_dependencies(mp: pytest.MonkeyPatch, block: PDFParserBlock, handler: AsyncMock) -> AsyncMock:
        workflow_run_context = MagicMock()
        workflow_run_context.has_parameter.return_value = False
        record_output_parameter_value = AsyncMock()

        async def fake_build_block_result(self: PDFParserBlock, **kwargs: Any) -> BlockResult:
            kwargs.pop("organization_id", None)
            return BlockResult(output_parameter=self.output_parameter, **kwargs)

        mp.setattr(
            PDFParserBlock, "get_workflow_run_context", staticmethod(lambda workflow_run_id: workflow_run_context)
        )
        mp.setattr(PDFParserBlock, "record_output_parameter_value", record_output_parameter_value)
        mp.setattr(PDFParserBlock, "build_block_result", fake_build_block_result)
        mp.setattr("skyvern.forge.sdk.api.files.download_file", AsyncMock(return_value="/tmp/test.pdf"))
        mp.setattr("skyvern.forge.sdk.workflow.models.block.extract_pdf_file", MagicMock(return_value="name\nAlice"))
        mp.setattr(
            "skyvern.forge.sdk.workflow.models.block.prompt_engine.load_prompt",
            MagicMock(return_value="base prompt"),
        )
        mp.setattr("skyvern.forge.sdk.workflow.models.block.app.LLM_API_HANDLER", handler)
        return record_output_parameter_value

    async def test_execute_retries_schema_validation_without_prevalidation_coercion(self) -> None:
        block = _make_pdf_parser_block("https://example.com/data.pdf")
        block.json_schema = {
            "type": "object",
            "properties": {"name": {"type": "string"}},
            "required": ["name"],
        }
        responses: list[Any] = [[{"name": "Alice"}], {"name": "Alice"}]
        prompts: list[str] = []
        force_dict_values: list[bool] = []

        async def fake_handler(**kwargs: Any) -> Any:
            prompts.append(kwargs["prompt"])
            force_dict_values.append(kwargs["force_dict"])
            return responses.pop(0)

        with pytest.MonkeyPatch.context() as mp:
            handler = AsyncMock(side_effect=fake_handler)
            record_output_parameter_value = self._patch_execute_dependencies(mp, block, handler)

            result = await block.execute("workflow-run", "workflow-run-block", organization_id="org-1")

        assert result.success is True
        assert result.status == BlockStatus.completed
        assert result.output_parameter_value == {"name": "Alice"}
        assert handler.await_count == 2
        assert force_dict_values == [False, False]
        assert "previous response failed JSON schema validation" in prompts[1]
        assert "expected type object, got array" in prompts[1]
        record_output_parameter_value.assert_awaited_once()


@pytest.mark.asyncio
class TestOcrPdfPages:
    """Tests for the per-page vision-LLM OCR path for scanned PDFs (SKY-10960)."""

    @staticmethod
    def _patch_handler(mp: pytest.MonkeyPatch, handler: AsyncMock) -> None:
        mp.setattr(
            "skyvern.forge.sdk.workflow.models.block.LLMAPIHandlerFactory.get_override_llm_api_handler",
            lambda *a, **kw: handler,
        )
        mp.setattr(
            "skyvern.forge.sdk.workflow.models.block.prompt_engine.load_prompt",
            MagicMock(return_value="prompt"),
        )

    async def test_each_page_transcribed_and_concatenated_in_order(self) -> None:
        """One LLM call per page; every page is concatenated in page order with markers."""
        block = _make_file_parser_block("https://example.com/scan.pdf", FileType.PDF)
        page_text = {
            b"img-1": "Cover sheet",
            b"img-2": "Demographics",
            b"img-3": "Provider: JORDAN SAMPLE MD",
        }

        async def fake_handler(**kwargs: Any) -> dict[str, str]:
            (image,) = kwargs["screenshots"]
            return {"extracted_text": page_text[image]}

        with pytest.MonkeyPatch.context() as mp:
            handler = AsyncMock(side_effect=fake_handler)
            self._patch_handler(mp, handler)

            result = await block._ocr_pdf_pages([b"img-1", b"img-2", b"img-3"])

        # One call per page — the multi-page document is never sent as a single collapsed call.
        assert handler.await_count == 3
        for call in handler.await_args_list:
            assert len(call.kwargs["screenshots"]) == 1
        # Late-page content survives the single-call collapse.
        assert "Provider: JORDAN SAMPLE MD" in result
        assert "Cover sheet" in result and "Demographics" in result
        assert result.index("--- Page 1 ---") < result.index("--- Page 2 ---") < result.index("--- Page 3 ---")

    async def test_invalid_ocr_response_retries_with_sanitized_prompt(self) -> None:
        """A missing extracted_text field is retried without echoing invalid content."""
        block = _make_file_parser_block("https://example.com/scan.pdf", FileType.PDF)
        secret_like_value = "customer-private-value-" + ("x" * 300)
        prompts: list[str] = []
        force_dict_values: list[bool] = []

        async def fake_handler(**kwargs: Any) -> dict[str, str]:
            prompts.append(kwargs["prompt"])
            force_dict_values.append(kwargs["force_dict"])
            if len(prompts) == 1:
                return {"wrong_field": secret_like_value}
            return {"extracted_text": "Recovered page text"}

        with pytest.MonkeyPatch.context() as mp:
            handler = AsyncMock(side_effect=fake_handler)
            self._patch_handler(mp, handler)

            result = await block._ocr_pdf_pages([b"img-1"])

        assert handler.await_count == 2
        assert force_dict_values == [False, False]
        assert "Recovered page text" in result
        assert "previous OCR response failed JSON validation" in prompts[1]
        assert "must include extracted_text as a string" in prompts[1]
        assert secret_like_value not in prompts[1]
        assert "customer-private-value" not in prompts[1]

    async def test_order_preserved_when_a_later_page_resolves_first(self) -> None:
        """Output stays in page order even if an earlier page's call finishes last."""
        block = _make_file_parser_block("https://example.com/scan.pdf", FileType.PDF)

        async def fake_handler(**kwargs: Any) -> dict[str, str]:
            (image,) = kwargs["screenshots"]
            if image == b"slow-1":
                await asyncio.sleep(0.05)
                return {"extracted_text": "first page body"}
            return {"extracted_text": "second page body"}

        with pytest.MonkeyPatch.context() as mp:
            handler = AsyncMock(side_effect=fake_handler)
            self._patch_handler(mp, handler)

            result = await block._ocr_pdf_pages([b"slow-1", b"fast-2"])

        assert result.index("first page body") < result.index("second page body")

    async def test_failed_page_is_skipped_not_fatal(self) -> None:
        """A per-page OCR failure is logged and skipped; the other pages still extract."""
        block = _make_file_parser_block("https://example.com/scan.pdf", FileType.PDF)

        async def fake_handler(**kwargs: Any) -> dict[str, str]:
            (image,) = kwargs["screenshots"]
            if image == b"bad":
                raise RuntimeError("vision model timeout")
            return {"extracted_text": f"text for {image.decode()}"}

        with pytest.MonkeyPatch.context() as mp:
            handler = AsyncMock(side_effect=fake_handler)
            self._patch_handler(mp, handler)

            result = await block._ocr_pdf_pages([b"good-1", b"bad", b"good-3"])

        assert handler.await_count == 3
        assert "text for good-1" in result and "text for good-3" in result
        assert "--- Page 2 ---" not in result

    async def test_all_pages_failed_raises(self) -> None:
        """A total OCR outage (every page errors) propagates instead of returning empty text."""
        block = _make_file_parser_block("https://example.com/scan.pdf", FileType.PDF)

        async def fake_handler(**kwargs: Any) -> dict[str, str]:
            raise RuntimeError("vision model outage")

        with pytest.MonkeyPatch.context() as mp:
            handler = AsyncMock(side_effect=fake_handler)
            self._patch_handler(mp, handler)

            with pytest.raises(RuntimeError, match="vision model outage"):
                await block._ocr_pdf_pages([b"p1", b"p2", b"p3"])

    async def test_empty_pages_contribute_nothing(self) -> None:
        """Pages that OCR to empty text add neither a marker nor content."""
        block = _make_file_parser_block("https://example.com/scan.pdf", FileType.PDF)

        async def fake_handler(**kwargs: Any) -> dict[str, str]:
            (image,) = kwargs["screenshots"]
            return {"extracted_text": "" if image == b"blank" else "real content"}

        with pytest.MonkeyPatch.context() as mp:
            handler = AsyncMock(side_effect=fake_handler)
            self._patch_handler(mp, handler)

            result = await block._ocr_pdf_pages([b"blank", b"page-2"])

        assert "--- Page 1 ---" not in result
        assert "--- Page 2 ---" in result and "real content" in result

    async def test_truncates_at_page_boundary_on_token_limit(self) -> None:
        """Concatenation stops at a page boundary once the token budget is exceeded."""
        block = _make_file_parser_block("https://example.com/scan.pdf", FileType.PDF)

        async def fake_handler(**kwargs: Any) -> dict[str, str]:
            (image,) = kwargs["screenshots"]
            return {"extracted_text": f"content {image.decode()}"}

        with pytest.MonkeyPatch.context() as mp:
            handler = AsyncMock(side_effect=fake_handler)
            self._patch_handler(mp, handler)
            # Each page chunk counts as 10 tokens; a 15-token budget admits only the first page.
            mp.setattr("skyvern.forge.sdk.workflow.models.block.count_tokens", lambda text: 10)
            mp.setattr("skyvern.forge.sdk.workflow.models.block.MAX_FILE_PARSE_INPUT_TOKENS", 15)

            result = await block._ocr_pdf_pages([b"page-1", b"page-2", b"page-3"])

        assert "--- Page 1 ---" in result
        assert "--- Page 2 ---" not in result and "--- Page 3 ---" not in result

    async def test_parse_pdf_file_routes_empty_text_to_per_page_ocr(self) -> None:
        """A scanned PDF (no extractable text layer) is routed through per-page OCR."""
        block = _make_file_parser_block("https://example.com/scan.pdf", FileType.PDF)

        async def fake_handler(**kwargs: Any) -> dict[str, str]:
            (image,) = kwargs["screenshots"]
            return {"extracted_text": f"page {image.decode()}"}

        with pytest.MonkeyPatch.context() as mp:
            handler = AsyncMock(side_effect=fake_handler)
            self._patch_handler(mp, handler)
            mp.setattr("skyvern.forge.sdk.workflow.models.block.extract_pdf_file", lambda *a, **kw: "")
            mp.setattr(
                "skyvern.forge.sdk.workflow.models.block.render_pdf_pages_as_images",
                lambda *a, **kw: [b"A", b"B"],
            )

            result = await block._parse_pdf_file("/tmp/scan.pdf")

        assert handler.await_count == 2
        assert "page A" in result and "page B" in result
