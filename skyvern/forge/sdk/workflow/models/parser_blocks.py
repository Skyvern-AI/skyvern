from __future__ import annotations

import asyncio
import codecs
import csv
import json
import os
import zipfile
from datetime import date, datetime, time
from pathlib import Path, PurePosixPath
from typing import Any, ClassVar, Literal
from urllib.parse import urlparse

import docx
import filetype
import pandas as pd
import structlog
from charset_normalizer import from_bytes

from skyvern.constants import MAX_FILE_PARSE_INPUT_TOKENS, MAX_PDF_OCR_PAGES, PDF_OCR_PAGE_CONCURRENCY
from skyvern.exceptions import PDFParsingError
from skyvern.forge import app
from skyvern.forge.prompts import prompt_engine
from skyvern.forge.sdk.api.files import (
    get_path_for_workflow_download_directory,
    resolve_local_or_download_file,
)
from skyvern.forge.sdk.api.llm.api_handler import LLMAPIHandler
from skyvern.forge.sdk.api.llm.api_handler_factory import LLMAPIHandlerFactory
from skyvern.forge.sdk.api.llm.exceptions import InvalidLLMResponseFormat, InvalidLLMResponseType
from skyvern.forge.sdk.api.llm.schema_validator import validate_schema
from skyvern.forge.sdk.core import skyvern_context
from skyvern.forge.sdk.experimentation.llm_prompt_config import get_llm_handler_for_prompt_type
from skyvern.forge.sdk.utils.pdf_parser import extract_pdf_file, render_pdf_pages_as_images, validate_pdf_file
from skyvern.forge.sdk.utils.sanitization import sanitize_postgres_text
from skyvern.forge.sdk.workflow.context_manager import WorkflowRunContext
from skyvern.forge.sdk.workflow.exceptions import InvalidFileType
from skyvern.forge.sdk.workflow.models.block import (
    SCHEMA_VALIDATION_MAX_ATTEMPTS,
    SCHEMA_VALIDATION_MAX_ERRORS,
    _build_schema_validation_retry_prompt,
    _default_structured_output_schema,
    _is_schema_configuration_failure,
    _json_type_name,
    _llm_response_format_failure_reason,
    _validate_response_against_json_schema,
    extract_file_url_from_block_output,
    sanitize_filename,
)
from skyvern.forge.sdk.workflow.models.block_base import Block
from skyvern.forge.sdk.workflow.models.parameter import PARAMETER_TYPE
from skyvern.schemas.workflows import BlockResult, BlockStatus, BlockType, FileType
from skyvern.utils.token_counter import count_tokens

LOG = structlog.get_logger()


class FileParserBlock(Block):
    # There is a mypy bug with Literal. Without the type: ignore, mypy will raise an error:
    # Parameter 1 of Literal[...] cannot be of type "Any"
    block_type: Literal[BlockType.FILE_URL_PARSER] = BlockType.FILE_URL_PARSER  # type: ignore

    # FileParserBlock CSV constants
    _CSV_SNIFF_LINES = 5
    _CSV_BINARY_PREFIX_BYTES = 4096
    _CSV_UTF_BOMS = (codecs.BOM_UTF16_LE, codecs.BOM_UTF16_BE, codecs.BOM_UTF32_LE, codecs.BOM_UTF32_BE)
    # Bounded cap for legitimate wide cells (JSON blobs, long descriptions); applied only while parsing.
    _MAX_CSV_FIELD_SIZE_BYTES = 10 * 1024 * 1024
    # ZIP extraction guards (zip-bomb protection; sizes from central-directory metadata).
    # ClassVar keeps these plain class attributes — without it pydantic wraps underscore
    # names in ModelPrivateAttr and class-level access breaks.
    _MAX_ZIP_ARCHIVE_BYTES: ClassVar[int] = 512 * 1024 * 1024
    _MAX_ZIP_ENTRIES: ClassVar[int] = 1000
    _MAX_ZIP_UNCOMPRESSED_BYTES: ClassVar[int] = 1024**3
    _ZIP_JUNK_DIRS: ClassVar[tuple[str, ...]] = ("__MACOSX",)
    _ZIP_JUNK_FILES: ClassVar[tuple[str, ...]] = (".DS_Store", "Thumbs.db")
    # Classic EOCD + max comment + ZIP64 locator + ZIP64 EOCD fixed part.
    _ZIP_EOCD_TAIL_BYTES: ClassVar[int] = 65_557 + 20 + 56

    file_url: str
    file_type: FileType = FileType.AUTO_DETECT
    json_schema: dict[str, Any] | None = None
    schema_validation_max_attempts: ClassVar[int] = SCHEMA_VALIDATION_MAX_ATTEMPTS
    ocr_validation_max_attempts: ClassVar[int] = SCHEMA_VALIDATION_MAX_ATTEMPTS

    def get_failure_error_codes(self) -> list[str]:
        return ["FILE_PARSER_ERROR"]

    def get_all_parameters(
        self,
        workflow_run_id: str,
    ) -> list[PARAMETER_TYPE]:
        workflow_run_context = self.get_workflow_run_context(workflow_run_id)
        if self.file_url and workflow_run_context.has_parameter(self.file_url):
            return [workflow_run_context.get_parameter(self.file_url)]
        return []

    def format_potential_template_parameters(self, workflow_run_context: WorkflowRunContext) -> None:
        self.file_url = self.format_block_parameter_template_from_workflow_run_context(
            self.file_url, workflow_run_context
        )

        self._apply_workflow_system_prompt(workflow_run_context)

    @staticmethod
    def _validate_ocr_llm_response(llm_response: Any) -> str | None:
        if not isinstance(llm_response, dict):
            return (
                f"OCR response must be a JSON object with extracted_text string; got {_json_type_name(llm_response)}."
            )
        if not isinstance(llm_response.get("extracted_text"), str):
            return (
                "OCR response must include extracted_text as a string; "
                f"got {_json_type_name(llm_response.get('extracted_text'))}."
            )
        return None

    @staticmethod
    def _build_ocr_validation_retry_prompt(prompt: str, failure_reason: str) -> str:
        return (
            f"{prompt}\n\n"
            "Your previous OCR response failed JSON validation.\n"
            f"Validation error: {failure_reason}\n\n"
            'Retry the task. Return only valid JSON with this exact shape: {"extracted_text": "..."} '
            "Do not include markdown, code fences, explanatory text, or extra fields."
        )

    @staticmethod
    def _validate_ai_response_against_json_schema(response: Any, json_schema: dict[str, Any]) -> str | None:
        return _validate_response_against_json_schema(
            response,
            json_schema,
            "File parser",
            max_errors=SCHEMA_VALIDATION_MAX_ERRORS,
        )

    def _detect_file_type_from_url(self, file_url: str, file_path: str | None = None) -> FileType:
        """Detect file type based on file extension in the URL, with magic-byte fallback."""
        url_parsed = urlparse(file_url)
        suffix = Path(url_parsed.path).suffix.lower()
        if suffix in (".xlsx", ".xls", ".xlsm"):
            return FileType.EXCEL
        elif suffix == ".pdf":
            return FileType.PDF
        elif suffix == ".tsv":
            return FileType.CSV  # TSV files are handled by the CSV parser
        elif suffix in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".tiff", ".tif"):
            return FileType.IMAGE
        elif suffix == ".docx":
            return FileType.DOCX
        elif suffix == ".doc":
            raise InvalidFileType(
                file_url=file_url,
                file_type=FileType.DOCX,
                error="Legacy .doc format (Word 97-2003) is not supported. Please convert the file to .docx format.",
            )
        elif suffix == ".zip":
            return FileType.ZIP
        elif suffix == ".csv":
            return FileType.CSV

        # URL extension is missing or unrecognized — try magic-byte detection on the downloaded file
        if file_path:
            detected = self._detect_file_type_from_magic_bytes(file_path)
            if detected is not None:
                LOG.info(
                    "FileParserBlock Detected file type from magic bytes (URL had no recognizable extension)",
                    file_url=file_url,
                    detected_file_type=detected,
                )
                return detected

        return FileType.CSV  # Final fallback for truly unknown files

    def _detect_file_type_from_magic_bytes(self, file_path: str) -> FileType | None:
        """Detect file type from magic bytes using the filetype library. Returns None if unrecognized."""
        kind = filetype.guess(file_path)
        if kind is None:
            return None

        mime = kind.mime
        if mime == "application/pdf":
            return FileType.PDF
        elif mime in (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel",
        ):
            return FileType.EXCEL
        elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return FileType.DOCX
        elif mime == "application/zip":
            # OOXML files are ZIP containers matched before this generic branch, so only plain archives reach it.
            return FileType.ZIP
        elif mime.startswith("image/"):
            return FileType.IMAGE
        return None

    def _detect_file_encoding(self, file_path: str) -> str:
        """Detect the encoding of a file using charset-normalizer with fallbacks.

        Reads a sample of the file (first 64KB) to detect encoding efficiently.
        Falls back through common encodings if detection fails.
        """
        sample_size = 65536  # 64KB sample for detection
        with open(file_path, "rb") as f:
            raw_data = f.read(sample_size)

        result = from_bytes(raw_data)
        best_match = result.best()
        if best_match and best_match.encoding:
            return best_match.encoding

        for encoding in ["utf-8", "cp1252", "latin-1"]:
            try:
                raw_data.decode(encoding)
                return encoding
            except UnicodeDecodeError:
                continue

        # latin-1 always succeeds (1:1 byte mapping), so this is a safety fallback
        return "latin-1"

    def _sniff_csv_delimiter(self, file_path: str) -> tuple[str, str]:
        """Return (delimiter, encoding). Samples full lines to avoid mid-row truncation."""
        # Read small raw byte prefix to quickly detect empty binary files before attempting text decoding/sniffing
        with open(file_path, "rb") as f:
            raw_prefix = f.read(self._CSV_BINARY_PREFIX_BYTES)
        # Reject files that contain no meaningful bytes
        if not raw_prefix.strip():
            raise csv.Error("File is empty")
        # Reject likely binary content:
        # - Presence of null bytes is a strong binary signal
        # - Exception: UTF-16/UTF-32 text often starts with BOM and may contain null bytes
        if b"\x00" in raw_prefix and not raw_prefix.startswith(self._CSV_UTF_BOMS):
            raise csv.Error("File contains binary data")

        # Detect best text encoding for file, then read only the first N full lines so csv.Sniffer sees complete rows
        encoding = self._detect_file_encoding(file_path)
        with open(file_path, encoding=encoding, errors="replace", newline="") as file:
            lines: list[str] = []
            for _ in range(self._CSV_SNIFF_LINES):
                line = file.readline()
                if not line:
                    break
                lines.append(line)

        # Build the sniffer sample from complete lines only
        sample = "".join(lines)
        # Guard against files that decode but still contain no meaningful text
        if not sample.strip():
            raise csv.Error("File is empty")

        try:
            delimiter = csv.Sniffer().sniff(sample).delimiter
        except csv.Error:
            delimiter = "\t" if file_path.lower().endswith(".tsv") else ","
        return delimiter, encoding

    def validate_file_type(self, file_url_used: str, file_path: str) -> None:
        if self.file_type == FileType.CSV:
            try:
                self._sniff_csv_delimiter(file_path)
            except csv.Error as e:
                raise InvalidFileType(file_url=file_url_used, file_type=self.file_type, error=str(e))
        elif self.file_type == FileType.EXCEL:
            try:
                # Try to read the file with pandas to validate it's a valid Excel file
                pd.read_excel(file_path, nrows=1, engine="calamine")
            except Exception as e:
                raise InvalidFileType(
                    file_url=file_url_used, file_type=self.file_type, error=f"Invalid Excel file format: {str(e)}"
                )
        elif self.file_type == FileType.PDF:
            try:
                validate_pdf_file(file_path, file_identifier=file_url_used)
            except PDFParsingError as e:
                raise InvalidFileType(file_url=file_url_used, file_type=self.file_type, error=str(e))
        elif self.file_type == FileType.IMAGE:
            kind = filetype.guess(file_path)
            if kind is None or not kind.mime.startswith("image/"):
                raise InvalidFileType(
                    file_url=file_url_used, file_type=self.file_type, error="File is not a valid image"
                )
        elif self.file_type == FileType.DOCX:
            try:
                # Try to open the file with python-docx to validate it's a valid DOCX file
                docx.Document(file_path)
            except Exception as e:
                raise InvalidFileType(
                    file_url=file_url_used, file_type=self.file_type, error=f"Invalid DOCX file format: {str(e)}"
                )
        elif self.file_type == FileType.ZIP:
            if not zipfile.is_zipfile(file_path):
                raise InvalidFileType(
                    file_url=file_url_used, file_type=self.file_type, error="File is not a valid ZIP archive"
                )

    async def _parse_csv_file(self, file_path: str) -> list[dict[str, Any]]:
        """Parse CSV/TSV file and return list of dictionaries."""
        delimiter, encoding = self._sniff_csv_delimiter(file_path)
        previous_limit = csv.field_size_limit(self._MAX_CSV_FIELD_SIZE_BYTES)
        try:
            with open(file_path, encoding=encoding, errors="replace", newline="") as file:
                reader = csv.DictReader(file, delimiter=delimiter)
                return list(reader)
        finally:
            csv.field_size_limit(previous_limit)

    def _clean_dataframe_for_json(self, df: pd.DataFrame) -> list[dict[str, Any]]:
        """Clean DataFrame to ensure it can be serialized to JSON."""
        # Replace NaN and NaT values with "nan" string
        df_cleaned = df.replace({pd.NA: "nan", pd.NaT: "nan"})
        df_cleaned = df_cleaned.where(pd.notna(df_cleaned), "nan")

        # Convert to list of dictionaries
        records = df_cleaned.to_dict("records")

        # Additional cleaning for any remaining problematic values
        for record in records:
            for key, value in record.items():
                if pd.isna(value) or value == "NaN" or value == "NaT":
                    record[key] = "nan"
                elif isinstance(value, (pd.Timestamp, datetime, date, time)):
                    # NaT timestamps are already caught by pd.isna() above, so this is always valid
                    record[key] = value.isoformat()
                elif isinstance(value, pd.Timedelta):
                    record[key] = str(value)

        return records

    async def _parse_excel_file(self, file_path: str) -> list[dict[str, Any]]:
        """Parse Excel file and return list of dictionaries."""
        try:
            # Read Excel file with pandas, specifying engine explicitly
            df = pd.read_excel(file_path, engine="calamine")
            # Clean and convert DataFrame to list of dictionaries
            return self._clean_dataframe_for_json(df)
        except ImportError as e:
            raise InvalidFileType(
                file_url=self.file_url,
                file_type=self.file_type,
                error=f"Missing required dependency for Excel parsing: {str(e)}. Please install calamine: pip install python-calamine",
            )
        except Exception as e:
            raise InvalidFileType(
                file_url=self.file_url, file_type=self.file_type, error=f"Failed to parse Excel file: {str(e)}"
            )

    async def _parse_pdf_file(
        self,
        file_path: str,
        workflow_run_block_id: str | None = None,
        organization_id: str | None = None,
    ) -> str:
        """Parse PDF file and return extracted text.

        Uses the shared PDF parsing utility that tries pypdf first,
        then falls back to pdfplumber if pypdf fails. If text extraction
        yields empty/minimal content (e.g. scanned or image-based PDFs),
        renders pages as images and sends them to a vision LLM for OCR.
        """
        try:
            extracted_text = extract_pdf_file(file_path, file_identifier=self.file_url)
        except PDFParsingError as e:
            raise InvalidFileType(file_url=self.file_url, file_type=self.file_type, error=str(e))

        # If text extraction returned meaningful content, use it directly
        if extracted_text.strip():
            return extracted_text

        # Scanned / image-based PDF — render pages as images and OCR each page in
        # its own vision-LLM call. A single call covering every page collapses a
        # multi-page document down to the first page or two.
        LOG.info(
            "PDF text extraction returned empty content, falling back to vision LLM OCR",
            file_url=self.file_url,
        )
        try:
            page_images = await asyncio.to_thread(
                render_pdf_pages_as_images,
                file_path,
                file_identifier=self.file_url,
                max_pages=MAX_PDF_OCR_PAGES,
            )
            if not page_images:
                return extracted_text
            return await self._ocr_pdf_pages(
                page_images,
                workflow_run_block_id=workflow_run_block_id,
                organization_id=organization_id,
            )
        except Exception:
            LOG.exception(
                "Failed to extract text from PDF via vision LLM fallback",
                file_url=self.file_url,
            )
            raise

    async def _resolve_file_parser_handler(
        self, prompt_type: str, distinct_id: str | None, organization_id: str | None
    ) -> LLMAPIHandler:
        """Resolve the default handler for a file-parser prompt type.

        Honors the LLM_CONFIG_BY_PROMPT_TYPE PostHog flag (keyed by prompt type) so the
        OCR and extraction models can be set without a deploy; falls back to the primary
        handler. A block-level override_llm_key still takes precedence at the call site.
        """
        if distinct_id:
            posthog_handler = await get_llm_handler_for_prompt_type(prompt_type, distinct_id, organization_id)
            if posthog_handler:
                return posthog_handler
        return app.LLM_API_HANDLER

    async def _ocr_pdf_pages(
        self,
        page_images: list[bytes],
        workflow_run_block_id: str | None = None,
        organization_id: str | None = None,
    ) -> str:
        """OCR each rendered PDF page in its own vision-LLM call and concatenate.

        Per-page transcription avoids the single-call collapse where a multi-page
        document is summarized down to its first page(s). Pages are transcribed with
        bounded concurrency, reassembled in page order with page markers, and
        truncated at a page boundary once MAX_FILE_PARSE_INPUT_TOKENS is reached.
        """
        if self.ocr_validation_max_attempts <= 0:
            raise ValueError("OCR validation max attempts must be greater than 0.")

        llm_prompt = prompt_engine.load_prompt("extract-text-from-image")
        default_handler = await self._resolve_file_parser_handler(
            "extract-text-from-image", workflow_run_block_id, organization_id
        )
        llm_api_handler = LLMAPIHandlerFactory.get_override_llm_api_handler(
            self.override_llm_key_for_organization(organization_id), default=default_handler
        )
        semaphore = asyncio.Semaphore(PDF_OCR_PAGE_CONCURRENCY)

        async def _ocr_page(page_image: bytes) -> str:
            async with semaphore:
                prompt_for_attempt = llm_prompt
                for attempt in range(self.ocr_validation_max_attempts):
                    try:
                        # OCR transcription intentionally skips system_prompt; it still applies
                        # to the downstream extract-information-from-file-text call.
                        llm_response = await llm_api_handler(
                            prompt=prompt_for_attempt,
                            prompt_name="extract-text-from-image",
                            screenshots=[page_image],
                            # Schema validation must inspect the raw parsed root; dict coercion can hide bad OCR JSON.
                            force_dict=False,
                            workflow_run_block_id=workflow_run_block_id,
                            organization_id=organization_id,
                        )
                    except (InvalidLLMResponseFormat, InvalidLLMResponseType) as e:
                        failure_reason = _llm_response_format_failure_reason(e)
                        will_retry = attempt + 1 < self.ocr_validation_max_attempts
                        LOG.warning(
                            "FileParserBlock PDF OCR LLM response failed response-format validation",
                            file_url=self.file_url,
                            attempt=attempt + 1,
                            max_attempts=self.ocr_validation_max_attempts,
                            will_retry=will_retry,
                            error_type=type(e).__name__,
                        )
                        if not will_retry:
                            raise ValueError(failure_reason) from e
                        prompt_for_attempt = self._build_ocr_validation_retry_prompt(llm_prompt, failure_reason)
                        continue

                    ocr_failure_reason = self._validate_ocr_llm_response(llm_response)
                    if not ocr_failure_reason:
                        return llm_response.get("extracted_text", "") or ""

                    will_retry = attempt + 1 < self.ocr_validation_max_attempts
                    LOG.warning(
                        "FileParserBlock PDF OCR LLM response failed schema validation",
                        file_url=self.file_url,
                        attempt=attempt + 1,
                        max_attempts=self.ocr_validation_max_attempts,
                        will_retry=will_retry,
                        failure_reason=ocr_failure_reason,
                    )
                    if not will_retry:
                        raise ValueError(ocr_failure_reason)
                    prompt_for_attempt = self._build_ocr_validation_retry_prompt(llm_prompt, ocr_failure_reason)
                raise RuntimeError("OCR retry loop exhausted without returning or raising.")

        page_results = await asyncio.gather(
            *(_ocr_page(page_image) for page_image in page_images),
            return_exceptions=True,
        )

        # A total OCR outage must fail the block, not record an empty success — match the
        # prior single-call path, which propagated OCR errors. Partial failures are skipped below.
        errors = [r for r in page_results if isinstance(r, BaseException)]
        if errors and len(errors) == len(page_results):
            raise errors[0]

        page_chunks: list[str] = []
        current_tokens = 0
        for page_number, result in enumerate(page_results, start=1):
            if isinstance(result, BaseException):
                LOG.warning(
                    "Failed to OCR a PDF page via vision LLM, skipping it",
                    file_url=self.file_url,
                    page_number=page_number,
                    error=str(result),
                )
                continue
            page_text = result.strip()
            if not page_text:
                continue
            chunk = f"--- Page {page_number} ---\n{page_text}"
            chunk_tokens = count_tokens(chunk)
            if current_tokens + chunk_tokens > MAX_FILE_PARSE_INPUT_TOKENS:
                LOG.warning(
                    "PDF OCR text exceeds token limit, truncating at page boundary",
                    file_url=self.file_url,
                    pages_included=page_number - 1,
                    total_pages=len(page_results),
                    max_tokens=MAX_FILE_PARSE_INPUT_TOKENS,
                )
                break
            current_tokens += chunk_tokens
            page_chunks.append(chunk)

        return "\n\n".join(page_chunks)

    async def _parse_image_file(
        self,
        file_path: str,
        workflow_run_block_id: str | None = None,
        organization_id: str | None = None,
    ) -> str:
        """Parse image file using vision LLM for OCR."""
        if self.ocr_validation_max_attempts <= 0:
            raise ValueError("OCR validation max attempts must be greater than 0.")

        try:
            with open(file_path, "rb") as f:
                image_bytes = f.read()

            llm_prompt = prompt_engine.load_prompt("extract-text-from-image")
            default_handler = await self._resolve_file_parser_handler(
                "extract-text-from-image", workflow_run_block_id, organization_id
            )
            llm_api_handler = LLMAPIHandlerFactory.get_override_llm_api_handler(
                self.override_llm_key_for_organization(organization_id), default=default_handler
            )
            # OCR transcription intentionally skips system_prompt — see
            # _parse_pdf_file_with_vision_ocr for rationale.
            prompt_for_attempt = llm_prompt
            for attempt in range(self.ocr_validation_max_attempts):
                try:
                    llm_response = await llm_api_handler(
                        prompt=prompt_for_attempt,
                        prompt_name="extract-text-from-image",
                        screenshots=[image_bytes],
                        # Schema validation must inspect the raw parsed root; dict coercion can hide bad OCR JSON.
                        force_dict=False,
                        workflow_run_block_id=workflow_run_block_id,
                        organization_id=organization_id,
                    )
                except (InvalidLLMResponseFormat, InvalidLLMResponseType) as e:
                    failure_reason = _llm_response_format_failure_reason(e)
                    will_retry = attempt + 1 < self.ocr_validation_max_attempts
                    LOG.warning(
                        "FileParserBlock image OCR LLM response failed response-format validation",
                        file_url=self.file_url,
                        attempt=attempt + 1,
                        max_attempts=self.ocr_validation_max_attempts,
                        will_retry=will_retry,
                        error_type=type(e).__name__,
                    )
                    if not will_retry:
                        raise ValueError(failure_reason) from e
                    prompt_for_attempt = self._build_ocr_validation_retry_prompt(llm_prompt, failure_reason)
                    continue

                ocr_failure_reason = self._validate_ocr_llm_response(llm_response)
                if not ocr_failure_reason:
                    return llm_response.get("extracted_text", "") or ""

                will_retry = attempt + 1 < self.ocr_validation_max_attempts
                LOG.warning(
                    "FileParserBlock image OCR LLM response failed schema validation",
                    file_url=self.file_url,
                    attempt=attempt + 1,
                    max_attempts=self.ocr_validation_max_attempts,
                    will_retry=will_retry,
                    failure_reason=ocr_failure_reason,
                )
                if not will_retry:
                    raise ValueError(ocr_failure_reason)
                prompt_for_attempt = self._build_ocr_validation_retry_prompt(llm_prompt, ocr_failure_reason)

            raise RuntimeError("OCR retry loop exhausted without returning or raising.")
        except Exception:
            LOG.exception("Failed to extract text from image via OCR", file_url=self.file_url)
            raise

    async def _parse_docx_file(self, file_path: str, max_tokens: int = MAX_FILE_PARSE_INPUT_TOKENS) -> str:
        """Parse DOCX file and return extracted text.

        Extracts text from all paragraphs and tables in the document,
        respecting the token limit.
        """
        try:
            document = docx.Document(file_path)
            text_parts = []
            current_tokens = 0
            truncated = False

            # Extract text from paragraphs
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    para_tokens = count_tokens(paragraph.text)
                    if max_tokens and current_tokens + para_tokens > max_tokens:
                        LOG.warning(
                            "DOCX text exceeds token limit, truncating",
                            file_url=self.file_url,
                            current_tokens=current_tokens,
                            max_tokens=max_tokens,
                        )
                        truncated = True
                        break
                    text_parts.append(paragraph.text)
                    current_tokens += para_tokens

            # Extract text from tables (only if not already truncated)
            if not truncated:
                for table in document.tables:
                    if truncated:
                        break
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            row_str = " | ".join(row_text)
                            row_tokens = count_tokens(row_str)
                            if max_tokens and current_tokens + row_tokens > max_tokens:
                                LOG.warning(
                                    "DOCX text exceeds token limit, truncating at table",
                                    file_url=self.file_url,
                                    current_tokens=current_tokens,
                                    max_tokens=max_tokens,
                                )
                                truncated = True
                                break
                            text_parts.append(row_str)
                            current_tokens += row_tokens

            extracted_text = "\n".join(text_parts)
            extracted_text = sanitize_postgres_text(extracted_text)
            LOG.info(
                "Successfully parsed DOCX file",
                file_url=self.file_url,
                paragraph_count=len(document.paragraphs),
                table_count=len(document.tables),
                text_length=len(extracted_text),
                truncated=truncated,
            )
            return extracted_text
        except Exception as e:
            raise InvalidFileType(
                file_url=self.file_url, file_type=self.file_type, error=f"Failed to parse DOCX file: {str(e)}"
            )

    @classmethod
    def _is_zip_junk_member(cls, member_name: str) -> bool:
        parts = PurePosixPath(member_name).parts
        if not parts:
            return True
        if any(part in cls._ZIP_JUNK_DIRS for part in parts):
            return True
        return parts[-1] in cls._ZIP_JUNK_FILES or parts[-1].startswith("._")

    @classmethod
    def _read_zip_total_entry_count(cls, file_path: str) -> int | None:
        try:
            file_size = os.path.getsize(file_path)
            tail_size = min(file_size, cls._ZIP_EOCD_TAIL_BYTES)
            with open(file_path, "rb") as file:
                file.seek(file_size - tail_size)
                tail = file.read(tail_size)

            eocd_index = tail.rfind(b"PK\x05\x06")
            if eocd_index < 0 or eocd_index + 22 > len(tail):
                return None

            count = int.from_bytes(tail[eocd_index + 10 : eocd_index + 12], "little")
            if count != 0xFFFF:
                return count

            zip64_eocd_index = tail.rfind(b"PK\x06\x06")
            if zip64_eocd_index < 0 or zip64_eocd_index + 40 > len(tail):
                return None
            return int.from_bytes(tail[zip64_eocd_index + 32 : zip64_eocd_index + 40], "little")
        except Exception:
            return None

    def _check_extracted_size_within_limit(self, total_bytes: int) -> None:
        if total_bytes > self._MAX_ZIP_UNCOMPRESSED_BYTES:
            raise InvalidFileType(
                file_url=self.file_url,
                file_type=self.file_type,
                error=f"ZIP archive uncompressed content exceeds the limit of {self._MAX_ZIP_UNCOMPRESSED_BYTES} bytes",
            )

    def _extract_zip_file(
        self, file_path: str, workflow_run_id: str, workflow_run_block_id: str
    ) -> list[dict[str, Any]]:
        """Extract a ZIP archive into the run's download directory.

        Returns the extracted files as {"file_name", "file_path", "file_size"} dicts sorted by
        file_name, so downstream blocks can consume the files from the local filesystem.
        """
        context = skyvern_context.current()
        run_id = context.run_id if context and context.run_id else workflow_run_id
        zip_stem = sanitize_filename(Path(file_path).stem, default="archive")
        extract_dir = (
            get_path_for_workflow_download_directory(run_id) / "unzipped" / f"{zip_stem}_{workflow_run_block_id}"
        )

        archive_size = os.path.getsize(file_path)
        if archive_size > self._MAX_ZIP_ARCHIVE_BYTES:
            raise InvalidFileType(
                file_url=self.file_url,
                file_type=self.file_type,
                error=f"ZIP archive size ({archive_size} bytes) exceeds the limit of {self._MAX_ZIP_ARCHIVE_BYTES} bytes",
            )

        declared_entry_count = self._read_zip_total_entry_count(file_path)
        if declared_entry_count is not None and declared_entry_count > self._MAX_ZIP_ENTRIES:
            raise InvalidFileType(
                file_url=self.file_url,
                file_type=self.file_type,
                error=f"ZIP archive declares {declared_entry_count} entries, exceeding the limit of {self._MAX_ZIP_ENTRIES}",
            )

        with zipfile.ZipFile(file_path) as zip_file:
            members = [
                member
                for member in zip_file.infolist()
                if not member.is_dir() and not self._is_zip_junk_member(member.filename)
            ]
            if len(members) > self._MAX_ZIP_ENTRIES:
                raise InvalidFileType(
                    file_url=self.file_url,
                    file_type=self.file_type,
                    error=f"ZIP archive contains {len(members)} files, exceeding the limit of {self._MAX_ZIP_ENTRIES}",
                )
            total_uncompressed_bytes = sum(member.file_size for member in members)
            # The declared-size check is advisory; measured bytes after extraction are authoritative.
            if total_uncompressed_bytes > self._MAX_ZIP_UNCOMPRESSED_BYTES:
                raise InvalidFileType(
                    file_url=self.file_url,
                    file_type=self.file_type,
                    error=f"ZIP archive uncompressed size ({total_uncompressed_bytes} bytes) exceeds the limit of {self._MAX_ZIP_UNCOMPRESSED_BYTES} bytes",
                )
            if any(member.flag_bits & 0x1 for member in members):
                raise InvalidFileType(
                    file_url=self.file_url,
                    file_type=self.file_type,
                    error="Password-protected ZIP archives are not supported",
                )

            extract_dir.mkdir(parents=True, exist_ok=True)
            # Keyed by destination path: member names that sanitize to the same destination
            # (e.g. "a.csv", "/a.csv", "../a.csv") overwrite on disk, and per ZIP semantics the
            # last entry wins — keep one list entry per file instead of duplicates.
            # measured_total_bytes intentionally counts every member's written bytes (including
            # overwritten collisions) because it guards total write I/O, not final disk usage.
            files_by_path: dict[str, dict[str, Any]] = {}
            measured_total_bytes = 0
            for member in members:
                if member.file_size > self._MAX_ZIP_UNCOMPRESSED_BYTES - measured_total_bytes:
                    raise InvalidFileType(
                        file_url=self.file_url,
                        file_type=self.file_type,
                        error=f"ZIP archive uncompressed content exceeds the limit of {self._MAX_ZIP_UNCOMPRESSED_BYTES} bytes",
                    )
                # ZipFile.extract sanitizes absolute paths and ".." components, so members cannot
                # escape extract_dir.
                extracted_path = zip_file.extract(member, path=extract_dir)
                extracted_size = Path(extracted_path).stat().st_size
                measured_total_bytes += extracted_size
                self._check_extracted_size_within_limit(measured_total_bytes)
                if extracted_path in files_by_path:
                    LOG.warning(
                        "FileParserBlock ZIP members collide after path sanitization, keeping the last one",
                        file_url=self.file_url,
                        member_name=member.filename,
                    )
                files_by_path[extracted_path] = {
                    "file_name": str(Path(extracted_path).relative_to(extract_dir)),
                    "file_path": extracted_path,
                    "file_size": extracted_size,
                }

        extracted_files = sorted(files_by_path.values(), key=lambda file_info: file_info["file_name"])
        LOG.info(
            "FileParserBlock Extracted ZIP archive",
            file_url=self.file_url,
            extract_dir=str(extract_dir),
            file_count=len(extracted_files),
        )
        return extracted_files

    async def _parse_file_of_type(
        self,
        file_type: FileType,
        file_path: str,
        workflow_run_block_id: str | None = None,
        organization_id: str | None = None,
    ) -> str | list[dict[str, Any]] | None:
        """Parse a file with the parser for its type; returns None for unsupported types."""
        if file_type == FileType.CSV:
            return await self._parse_csv_file(file_path)
        if file_type == FileType.EXCEL:
            return await self._parse_excel_file(file_path)
        if file_type == FileType.PDF:
            return await self._parse_pdf_file(
                file_path, workflow_run_block_id=workflow_run_block_id, organization_id=organization_id
            )
        if file_type == FileType.IMAGE:
            return await self._parse_image_file(
                file_path, workflow_run_block_id=workflow_run_block_id, organization_id=organization_id
            )
        if file_type == FileType.DOCX:
            return await self._parse_docx_file(file_path)
        return None

    async def _parse_zip_contents(
        self,
        extracted_files: list[dict[str, Any]],
        workflow_run_block_id: str | None = None,
        organization_id: str | None = None,
    ) -> list[dict[str, Any]]:
        """Parse every supported file extracted from a ZIP for combined AI extraction.

        Returns [{"file_name", "content"}, ...]. Unsupported or unparseable files are skipped;
        total content is truncated at a file boundary once MAX_FILE_PARSE_INPUT_TOKENS is reached.
        """
        content_entries: list[dict[str, Any]] = []
        current_tokens = 0
        for file_info in extracted_files:
            inner_path: str = file_info["file_path"]
            file_name: str = file_info["file_name"]
            try:
                inner_file_type = self._detect_file_type_from_url(inner_path, file_path=inner_path)
            except InvalidFileType as e:
                LOG.warning("FileParserBlock Skipping unsupported file in ZIP", file_name=file_name, error=str(e))
                continue
            if inner_file_type == FileType.ZIP:
                LOG.warning("FileParserBlock Skipping nested ZIP archive", file_name=file_name)
                continue
            try:
                content = await self._parse_file_of_type(
                    inner_file_type,
                    inner_path,
                    workflow_run_block_id=workflow_run_block_id,
                    organization_id=organization_id,
                )
            except Exception as e:
                LOG.warning(
                    "FileParserBlock Failed to parse file in ZIP, skipping it", file_name=file_name, error=str(e)
                )
                continue
            if content is None:
                LOG.warning(
                    "FileParserBlock Skipping file with unsupported type in ZIP",
                    file_name=file_name,
                    detected_file_type=inner_file_type,
                )
                continue
            entry = {"file_name": file_name, "content": content}
            entry_tokens = count_tokens(json.dumps(entry, separators=(",", ":"), default=str))
            if current_tokens + entry_tokens > MAX_FILE_PARSE_INPUT_TOKENS:
                if not content_entries:
                    raise InvalidFileType(
                        file_url=self.file_url,
                        file_type=self.file_type,
                        error=f"File '{file_name}' in the ZIP archive alone exceeds the maximum extraction input size",
                    )
                LOG.warning(
                    "FileParserBlock ZIP content exceeds token limit, truncating at file boundary",
                    file_url=self.file_url,
                    files_included=len(content_entries),
                    total_files=len(extracted_files),
                    max_tokens=MAX_FILE_PARSE_INPUT_TOKENS,
                )
                break
            current_tokens += entry_tokens
            content_entries.append(entry)

        if not content_entries:
            raise InvalidFileType(
                file_url=self.file_url,
                file_type=self.file_type,
                error="ZIP archive contains no parseable files (supported: CSV, Excel, PDF, image, DOCX)",
            )
        return content_entries

    async def _extract_with_ai(
        self,
        content: str | list[dict[str, Any]],
        workflow_run_context: WorkflowRunContext,
        workflow_run_block_id: str | None = None,
        organization_id: str | None = None,
    ) -> dict[str, Any] | list | str | None:
        """Extract structured data using AI based on json_schema."""
        # Use local variable to avoid mutating the instance
        schema_to_use = self.json_schema or _default_structured_output_schema("Information extracted from the file")
        if not validate_schema(schema_to_use):
            raise ValueError("File parser JSON schema is invalid.")

        # Convert content to string for AI processing
        if isinstance(content, list):
            content_str = json.dumps(content, separators=(",", ":"))
        else:
            content_str = content

        llm_prompt = prompt_engine.load_prompt(
            "extract-information-from-file-text", extracted_text_content=content_str, json_schema=schema_to_use
        )

        llm_key = self.override_llm_key_for_organization(organization_id)
        default_handler = await self._resolve_file_parser_handler(
            "extract-information-from-file-text", workflow_run_block_id, organization_id
        )
        llm_api_handler = LLMAPIHandlerFactory.get_override_llm_api_handler(llm_key, default=default_handler)

        prompt_for_attempt = llm_prompt
        for attempt in range(self.schema_validation_max_attempts):
            try:
                llm_response = await llm_api_handler(
                    prompt=prompt_for_attempt,
                    prompt_name="extract-information-from-file-text",
                    # Schema validation must inspect the raw parsed root; dict coercion can hide wrong-root responses.
                    force_dict=False,
                    system_prompt=self.workflow_system_prompt,
                    workflow_run_block_id=workflow_run_block_id,
                    organization_id=organization_id,
                )
            except (InvalidLLMResponseFormat, InvalidLLMResponseType) as e:
                failure_reason = _llm_response_format_failure_reason(e)
                will_retry = attempt + 1 < self.schema_validation_max_attempts
                LOG.warning(
                    "FileParserBlock extraction LLM response failed response-format validation",
                    file_url=self.file_url,
                    attempt=attempt + 1,
                    max_attempts=self.schema_validation_max_attempts,
                    will_retry=will_retry,
                    error_type=type(e).__name__,
                    schema_type=schema_to_use.get("type"),
                )
                if not will_retry:
                    raise ValueError(failure_reason) from e
                prompt_for_attempt = _build_schema_validation_retry_prompt(llm_prompt, failure_reason)
                continue

            schema_validation_failure = self._validate_ai_response_against_json_schema(llm_response, schema_to_use)
            if not schema_validation_failure:
                return llm_response

            is_schema_configuration_failure = _is_schema_configuration_failure(schema_validation_failure)
            will_retry = attempt + 1 < self.schema_validation_max_attempts and not is_schema_configuration_failure
            LOG.warning(
                "FileParserBlock extraction LLM response failed schema validation",
                file_url=self.file_url,
                attempt=attempt + 1,
                max_attempts=self.schema_validation_max_attempts,
                will_retry=will_retry,
                failure_reason=schema_validation_failure,
                schema_type=schema_to_use.get("type"),
            )
            if not will_retry:
                raise ValueError(schema_validation_failure)
            prompt_for_attempt = _build_schema_validation_retry_prompt(
                llm_prompt,
                schema_validation_failure,
            )

        raise AssertionError("unreachable schema validation retry loop exit")

    async def _record_failure(
        self,
        workflow_run_context: WorkflowRunContext,
        workflow_run_id: str,
        workflow_run_block_id: str,
        organization_id: str | None,
        failure_reason: str,
    ) -> BlockResult:
        # SKY-7939: also surface the failure in `outputs.<block_label>` so callers
        # can tell which block failed without cross-referencing the timeline.
        error_codes = self.get_failure_error_codes()
        failure_output: dict[str, Any] = {
            "status": BlockStatus.failed.value,
            "failure_reason": failure_reason,
            "errors": [
                {"error_code": code, "reasoning": failure_reason, "confidence_float": 1.0} for code in error_codes
            ],
        }
        await self.record_output_parameter_value(workflow_run_context, workflow_run_id, failure_output)
        return await self.build_block_result(
            success=False,
            failure_reason=failure_reason,
            output_parameter_value=failure_output,
            status=BlockStatus.failed,
            workflow_run_block_id=workflow_run_block_id,
            organization_id=organization_id,
            error_codes=error_codes or None,
        )

    @staticmethod
    def _extract_file_url_from_block_output(value: Any) -> str | None:
        """Extract a file URL from a block output value.

        When users pass an entire block output (e.g. ``{{ block_8_output }}``) as the
        ``file_url``, the resolved value may be a dict or a string representation of a
        dict that contains a ``downloaded_files`` list.  This helper unwraps that
        structure and returns the URL of the first downloaded file.

        Handles three forms:
        - dict with a ``downloaded_files`` list
        - JSON string encoding such a dict
        - Python dict-repr string produced by Jinja's default ``str()`` rendering
        """
        return extract_file_url_from_block_output(value)

    async def execute(
        self,
        workflow_run_id: str,
        workflow_run_block_id: str,
        organization_id: str | None = None,
        browser_session_id: str | None = None,
        **kwargs: dict,
    ) -> BlockResult:
        workflow_run_context = self.get_workflow_run_context(workflow_run_id)

        if (
            self.file_url
            and workflow_run_context.has_parameter(self.file_url)
            and workflow_run_context.has_value(self.file_url)
        ):
            file_url_parameter_value = workflow_run_context.get_value(self.file_url)
            if file_url_parameter_value:
                extracted_url = self._extract_file_url_from_block_output(file_url_parameter_value)
                if extracted_url:
                    LOG.info(
                        "FileParserBlock Extracted file URL from block output parameter",
                        extracted_url=extracted_url,
                        file_url_parameter_key=self.file_url,
                    )
                    self.file_url = extracted_url
                else:
                    LOG.info(
                        "FileParserBlock File URL is parameterized, using parameter value",
                        file_url_parameter_value=file_url_parameter_value,
                        file_url_parameter_key=self.file_url,
                    )
                    self.file_url = file_url_parameter_value

        try:
            self.format_potential_template_parameters(workflow_run_context)
        except Exception as e:
            return await self._record_failure(
                workflow_run_context,
                workflow_run_id,
                workflow_run_block_id,
                organization_id,
                f"Failed to format jinja template: {str(e)}",
            )

        # After Jinja rendering, self.file_url may be a stringified block output
        # (e.g. when the user wrote ``{{ block_8_output }}``). Try to extract the
        # file URL from it before attempting the download.
        extracted_url = self._extract_file_url_from_block_output(self.file_url)
        if extracted_url:
            LOG.info(
                "FileParserBlock Extracted file URL from rendered block output",
                extracted_url=extracted_url,
                rendered_value=self.file_url,
            )
            self.file_url = extracted_url

        try:
            context = skyvern_context.current()
            run_id = context.run_id if context and context.run_id else workflow_run_id
            file_path = await resolve_local_or_download_file(self.file_url, run_id, organization_id=organization_id)

            # Resolve AUTO_DETECT (and legacy CSV-as-default) via URL/magic-byte detection;
            # IMAGE/EXCEL/PDF/DOCX/ZIP are honored as user overrides.
            if self.file_type not in (FileType.IMAGE, FileType.EXCEL, FileType.PDF, FileType.DOCX, FileType.ZIP):
                self.file_type = self._detect_file_type_from_url(self.file_url, file_path=file_path)

            # Validate the file type
            self.validate_file_type(self.file_url, file_path)
        except Exception as e:
            return await self._record_failure(
                workflow_run_context,
                workflow_run_id,
                workflow_run_block_id,
                organization_id,
                f"Failed to download or validate file: {str(e)}",
            )

        LOG.debug(
            "FileParserBlock After file type validation",
            file_type=self.file_type,
            json_schema_present=self.json_schema is not None,
            json_schema_type=type(self.json_schema),
        )

        # Parse the file based on type
        parsed_data: str | list[dict[str, Any]]
        try:
            if self.file_type == FileType.ZIP:
                extracted_zip_files = await asyncio.to_thread(
                    self._extract_zip_file, file_path, workflow_run_id, workflow_run_block_id
                )
                if self.json_schema:
                    parsed_data = await self._parse_zip_contents(
                        extracted_zip_files,
                        workflow_run_block_id=workflow_run_block_id,
                        organization_id=organization_id,
                    )
                else:
                    parsed_data = extracted_zip_files
            else:
                maybe_parsed = await self._parse_file_of_type(
                    self.file_type,
                    file_path,
                    workflow_run_block_id=workflow_run_block_id,
                    organization_id=organization_id,
                )
                if maybe_parsed is None:
                    return await self._record_failure(
                        workflow_run_context,
                        workflow_run_id,
                        workflow_run_block_id,
                        organization_id,
                        f"Unsupported file type: {self.file_type}",
                    )
                parsed_data = maybe_parsed
        except Exception as e:
            return await self._record_failure(
                workflow_run_context,
                workflow_run_id,
                workflow_run_block_id,
                organization_id,
                f"Failed to parse {self.file_type} file: {str(e)}",
            )

        # If json_schema is provided, use AI to extract structured data
        final_data: Any
        LOG.debug(
            "FileParserBlock JSON schema check",
            has_json_schema=self.json_schema is not None,
            json_schema_type=type(self.json_schema),
            json_schema=self.json_schema,
        )

        if self.json_schema:
            try:
                ai_extracted_data = await self._extract_with_ai(
                    parsed_data,
                    workflow_run_context,
                    workflow_run_block_id=workflow_run_block_id,
                    organization_id=organization_id,
                )
                final_data = ai_extracted_data
            except Exception as e:
                return await self._record_failure(
                    workflow_run_context,
                    workflow_run_id,
                    workflow_run_block_id,
                    organization_id,
                    f"Failed to extract data with AI: {str(e)}",
                )
        else:
            # Return raw parsed data
            final_data = parsed_data

        # Record the parsed data
        await self.record_output_parameter_value(workflow_run_context, workflow_run_id, final_data)
        return await self.build_block_result(
            success=True,
            failure_reason=None,
            output_parameter_value=final_data,
            status=BlockStatus.completed,
            workflow_run_block_id=workflow_run_block_id,
            organization_id=organization_id,
        )


class PDFParserBlock(Block):
    """
    DEPRECATED: Use FileParserBlock with file_type=FileType.PDF instead.
    This block will be removed in a future version.
    """

    # There is a mypy bug with Literal. Without the type: ignore, mypy will raise an error:
    # Parameter 1 of Literal[...] cannot be of type "Any"
    block_type: Literal[BlockType.PDF_PARSER] = BlockType.PDF_PARSER  # type: ignore

    file_url: str
    json_schema: dict[str, Any] | None = None
    schema_validation_max_attempts: ClassVar[int] = SCHEMA_VALIDATION_MAX_ATTEMPTS

    def get_all_parameters(
        self,
        workflow_run_id: str,
    ) -> list[PARAMETER_TYPE]:
        workflow_run_context = self.get_workflow_run_context(workflow_run_id)
        if self.file_url and workflow_run_context.has_parameter(self.file_url):
            return [workflow_run_context.get_parameter(self.file_url)]
        return []

    def format_potential_template_parameters(self, workflow_run_context: WorkflowRunContext) -> None:
        self.file_url = self.format_block_parameter_template_from_workflow_run_context(
            self.file_url, workflow_run_context
        )

        self._apply_workflow_system_prompt(workflow_run_context)

    async def execute(
        self,
        workflow_run_id: str,
        workflow_run_block_id: str,
        organization_id: str | None = None,
        browser_session_id: str | None = None,
        **kwargs: dict,
    ) -> BlockResult:
        workflow_run_context = self.get_workflow_run_context(workflow_run_id)
        if (
            self.file_url
            and workflow_run_context.has_parameter(self.file_url)
            and workflow_run_context.has_value(self.file_url)
        ):
            file_url_parameter_value = workflow_run_context.get_value(self.file_url)
            if file_url_parameter_value:
                LOG.info(
                    "PDFParserBlock File URL is parameterized, using parameter value",
                    file_url_parameter_value=file_url_parameter_value,
                    file_url_parameter_key=self.file_url,
                )
                self.file_url = file_url_parameter_value

        try:
            self.format_potential_template_parameters(workflow_run_context)
        except Exception as e:
            return await self.build_block_result(
                success=False,
                failure_reason=f"Failed to format jinja template: {str(e)}",
                output_parameter_value=None,
                status=BlockStatus.failed,
                workflow_run_block_id=workflow_run_block_id,
                organization_id=organization_id,
            )

        try:
            context = skyvern_context.current()
            run_id = context.run_id if context and context.run_id else workflow_run_id
            file_path = await resolve_local_or_download_file(self.file_url, run_id, organization_id=organization_id)
        except Exception as e:
            return await self.build_block_result(
                success=False,
                failure_reason=f"Failed to download or validate file: {str(e)}",
                output_parameter_value=None,
                status=BlockStatus.failed,
                workflow_run_block_id=workflow_run_block_id,
                organization_id=organization_id,
            )

        try:
            extracted_text = extract_pdf_file(file_path, file_identifier=self.file_url)
        except PDFParsingError:
            return await self.build_block_result(
                success=False,
                failure_reason="Failed to parse PDF file",
                output_parameter_value=None,
                status=BlockStatus.failed,
                workflow_run_block_id=workflow_run_block_id,
                organization_id=organization_id,
            )

        if not self.json_schema:
            self.json_schema = _default_structured_output_schema("Information extracted from the text")
        schema_to_use = self.json_schema
        assert schema_to_use is not None
        if not validate_schema(schema_to_use):
            return await self.build_block_result(
                success=False,
                failure_reason="File parser JSON schema is invalid.",
                output_parameter_value=None,
                status=BlockStatus.failed,
                workflow_run_block_id=workflow_run_block_id,
                organization_id=organization_id,
            )

        llm_prompt = prompt_engine.load_prompt(
            "extract-information-from-file-text", extracted_text_content=extracted_text, json_schema=schema_to_use
        )

        llm_response: dict[str, Any] | list | str | None = None
        prompt_for_attempt = llm_prompt
        for attempt in range(self.schema_validation_max_attempts):
            try:
                llm_response = await app.LLM_API_HANDLER(
                    prompt=prompt_for_attempt,
                    prompt_name="extract-information-from-file-text",
                    # Schema validation must inspect the raw parsed root; dict coercion can hide wrong-root responses.
                    force_dict=False,
                    system_prompt=self.workflow_system_prompt,
                    workflow_run_block_id=workflow_run_block_id,
                    organization_id=organization_id,
                )
            except (InvalidLLMResponseFormat, InvalidLLMResponseType) as e:
                failure_reason = _llm_response_format_failure_reason(e)
                will_retry = attempt + 1 < self.schema_validation_max_attempts
                LOG.warning(
                    "PDFParserBlock extraction LLM response failed response-format validation",
                    file_url=self.file_url,
                    attempt=attempt + 1,
                    max_attempts=self.schema_validation_max_attempts,
                    will_retry=will_retry,
                    error_type=type(e).__name__,
                    schema_type=schema_to_use.get("type"),
                )
                if not will_retry:
                    return await self.build_block_result(
                        success=False,
                        failure_reason=failure_reason,
                        output_parameter_value=None,
                        status=BlockStatus.failed,
                        workflow_run_block_id=workflow_run_block_id,
                        organization_id=organization_id,
                    )
                prompt_for_attempt = _build_schema_validation_retry_prompt(llm_prompt, failure_reason)
                continue

            schema_validation_failure = FileParserBlock._validate_ai_response_against_json_schema(
                llm_response,
                schema_to_use,
            )
            if not schema_validation_failure:
                break

            is_schema_configuration_failure = _is_schema_configuration_failure(schema_validation_failure)
            will_retry = attempt + 1 < self.schema_validation_max_attempts and not is_schema_configuration_failure
            LOG.warning(
                "PDFParserBlock extraction LLM response failed schema validation",
                file_url=self.file_url,
                attempt=attempt + 1,
                max_attempts=self.schema_validation_max_attempts,
                will_retry=will_retry,
                failure_reason=schema_validation_failure,
                schema_type=schema_to_use.get("type"),
            )
            if not will_retry:
                return await self.build_block_result(
                    success=False,
                    failure_reason=schema_validation_failure,
                    output_parameter_value=None,
                    status=BlockStatus.failed,
                    workflow_run_block_id=workflow_run_block_id,
                    organization_id=organization_id,
                )
            prompt_for_attempt = _build_schema_validation_retry_prompt(
                llm_prompt,
                schema_validation_failure,
            )

        # Record the parsed data
        await self.record_output_parameter_value(workflow_run_context, workflow_run_id, llm_response)
        return await self.build_block_result(
            success=True,
            failure_reason=None,
            output_parameter_value=llm_response,
            status=BlockStatus.completed,
            workflow_run_block_id=workflow_run_block_id,
            organization_id=organization_id,
        )
